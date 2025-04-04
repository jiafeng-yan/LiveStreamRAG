import os
from typing import List, Dict, Any
import pypdf
import docx
import markdown
from langchain.docstore.document import Document
import re
import markdown
from markdown.extensions.extra import ExtraExtension
from bs4 import BeautifulSoup

class DocumentLoader:
    """处理各种文档格式的加载器"""

    @staticmethod
    def load(file_path: str) -> List[Document]:
        """根据文件类型加载文档"""
        _, extension = os.path.splitext(file_path.lower())

        if extension == ".pdf":
            return DocumentLoader._load_pdf(file_path)
        elif extension == ".docx":
            return DocumentLoader._load_docx(file_path)
        elif extension == ".md":
            return DocumentLoader._load_markdown(file_path)
        elif extension == ".txt":
            return DocumentLoader._load_text(file_path)
        else:
            raise ValueError(f"不支持的文档类型: {extension}")

    @staticmethod
    def _load_pdf(file_path: str) -> List[Document]:
        """加载PDF文档"""
        documents = []
        with open(file_path, "rb") as file:
            pdf = pypdf.PdfReader(file)
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text.strip():  # 跳过空白页
                    metadata = {"source": file_path, "page": i + 1}
                    documents.append(Document(page_content=text, metadata=metadata))
        return documents

    @staticmethod
    def _load_docx(file_path: str) -> List[Document]:
        """加载Word文档"""
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        metadata = {"source": file_path}
        return [Document(page_content=text, metadata=metadata)]

    @staticmethod
    def _load_markdown(file_path: str) -> List[Document]:
        """加载Markdown文档"""
        with open(file_path, "r", encoding="utf-8") as file:
            text = file.read()
        # 转换为纯文本 (去除标记)
        html = markdown.Markdown(extensions=[ExtraExtension()]).convert(text)
        # html = markdown.markdown(text)
        # 简单去除HTML标签 (实际应用中可能需要更复杂的处理)
        content = BeautifulSoup(html, "html.parser").get_text()
        # content = html.replace("<p>", "\n").replace("</p>", "\n")
        # content = re.sub(r"<[^>]*>", "", content)
        # 去除多余空行和首尾空格
        content = '\n'.join(
            line.strip() for line in content.splitlines() if line.strip()
        )
        metadata = {"source": file_path}
        return [Document(page_content=content, metadata=metadata)]

    @staticmethod
    def _load_text(file_path: str) -> List[Document]:
        """加载纯文本文档"""
        with open(file_path, "r", encoding="utf-8") as file:
            text = file.read()
        metadata = {"source": file_path}
        return [Document(page_content=text, metadata=metadata)]
